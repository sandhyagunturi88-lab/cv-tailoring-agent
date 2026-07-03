"""
core/report.py — deterministic fit scoring and export. No LLM calls: the
score shown on screen must be reproducible for the same hits/gaps/keywords.
"""

from __future__ import annotations

import io


def fit_score(hits: dict, gaps: list, keywords: list) -> int:
    weight = lambda k: 3 if k["priority"] == "must" else 1  # noqa: E731
    total = sum(weight(k) for k in keywords) or 1
    got = sum(weight(k) for kws in hits.values() for k in kws)
    return round(100 * got / total)


def export_docx(bullets: list[str]) -> bytes:
    try:
        from docx import Document  # type: ignore

        doc = Document()
        doc.add_heading("Tailored CV bullets", level=1)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return ("# Tailored CV bullets\n" + "\n".join(f"- {b}" for b in bullets)).encode()
